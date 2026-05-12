# backend/ml/exporter.py
# ─────────────────────────────────────────────────────────────
# Notes Export — TXT, DOCX, PDF
# All crashes from empty-paragraph .runs[0] fixed.
# ─────────────────────────────────────────────────────────────

import io
import re
from typing import Optional


def _strip_md(text: str) -> str:
    """Remove markdown symbols so output is clean plain text."""
    if not text:
        return ''
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',    r'\1', text)
    text = re.sub(r'`{1,3}([^`]+)`{1,3}', r'\1', text)
    text = re.sub(r'^[-*+]\s+', '• ', text, flags=re.MULTILINE)
    return text.strip()


def _is_bullet(line: str) -> bool:
    return line.startswith('• ') or re.match(r'^\d+[.)]\s', line) is not None


# ─────────────────────────────────────────────────────────────
#  TXT Export
# ─────────────────────────────────────────────────────────────

def export_txt(notes_dict: dict) -> bytes:
    from ml.note_structurer import NoteStructurer
    text = NoteStructurer.to_plain_text(notes_dict)
    return text.encode("utf-8")


# ─────────────────────────────────────────────────────────────
#  DOCX Export — properly structured with indentation
# ─────────────────────────────────────────────────────────────

def export_docx(notes_dict: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Helper: add paragraph with optional indent ────────────
    def body_para(text, indent_cm=0, bold=False, italic=False, color=None):
        p = doc.add_paragraph()
        if indent_cm:
            p.paragraph_format.left_indent = Cm(indent_cm)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def bullet_para(text, indent_cm=0.75):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.add_run(text)
        return p

    def num_para(text, indent_cm=0.75):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.add_run(text)
        return p

    # ── Title ─────────────────────────────────────────────────
    title_para = doc.add_heading(notes_dict.get("title", "Lecture Notes"), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Summary/Overview ──────────────────────────────────────
    summary = _strip_md(notes_dict.get("summary", ""))
    if summary:
        doc.add_heading("Overview", level=1)
        body_para(summary, italic=True, color=(107, 100, 88))

    doc.add_paragraph()

    # ── Key Points — numbered, indented ───────────────────────
    key_points = notes_dict.get("key_points", [])
    if key_points:
        doc.add_heading("Key Concepts", level=1)
        for pt in key_points:
            clean = _strip_md(pt)
            if clean:
                num_para(clean)

    doc.add_paragraph()

    # ── Detailed Sections ─────────────────────────────────────
    sections = notes_dict.get("sections", [])
    if sections:
        doc.add_heading("Detailed Notes", level=1)
        for sec in sections:
            heading = _strip_md(sec.get("heading", ""))
            if heading:
                doc.add_heading(heading, level=2)
            content = _strip_md(sec.get("content", ""))
            for raw_line in content.split('\n'):
                line = raw_line.strip()
                if not line:
                    continue
                if raw_line.startswith('  - ') or raw_line.startswith('  • '):
                    # Sub-bullet — deeper indent, muted colour
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Cm(2.0)
                    run = p.add_run(line.lstrip('-• '))
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 100)
                elif re.match(r'^  \d+[.)]', raw_line):
                    # Sub-numbered list — deeper indent, muted
                    p = doc.add_paragraph(style='List Number')
                    p.paragraph_format.left_indent = Cm(2.0)
                    run = p.add_run(re.sub(r'^\d+[.)] *', '', line))
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 100)
                elif line.startswith('• '):
                    bullet_para(line[2:], indent_cm=1.0)
                elif re.match(r'^\d+[.)]\s', line):
                    cleaned = re.sub(r'^\d+[.)]\s+', '', line)
                    num_para(cleaned, indent_cm=1.0)
                else:
                    body_para(line)

    doc.add_paragraph()

    # ── Full Transcript (if present) ─────────────────────────
    full_tr = notes_dict.get("full_transcript", "")
    if full_tr:
        doc.add_heading("Full Transcript (English)", level=1)
        for chunk in (full_tr[i:i+400] for i in range(0, len(full_tr), 400)):
            body_para(chunk, color=(107, 100, 88))

    # ── Footer ────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph("Generated by AudioNotes AI")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.italic = True
    for r in fp.runs:
        r.font.color.rgb = RGBColor(0xAA, 0xA5, 0x9E)
        r.font.size      = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────
#  PDF Export — structured, indented, clean
# ─────────────────────────────────────────────────────────────

def export_pdf(notes_dict: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
            ListFlowable, ListItem,
        )
    except ImportError:
        raise ImportError("reportlab is required: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2.2*cm, leftMargin=2.2*cm,
        topMargin=2.2*cm,  bottomMargin=2.2*cm,
        title=notes_dict.get("title", "Lecture Notes"),
    )

    styles = getSampleStyleSheet()
    BRAND  = colors.HexColor("#1F2937")   # dark slate
    ACCENT = colors.HexColor("#374151")
    MUTED  = colors.HexColor("#6B7280")
    RULE   = colors.HexColor("#E5E7EB")

    title_s = ParagraphStyle("NoteTitle", parent=styles["Title"],
        fontName="Helvetica-Bold", fontSize=20, textColor=BRAND,
        spaceAfter=4, alignment=1)

    h1_s = ParagraphStyle("H1", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=13, textColor=BRAND,
        spaceBefore=14, spaceAfter=5, borderPadding=0)

    h2_s = ParagraphStyle("H2", parent=styles["Heading2"],
        fontName="Helvetica-Bold", fontSize=11, textColor=ACCENT,
        spaceBefore=10, spaceAfter=3, leftIndent=10)

    body_s = ParagraphStyle("Body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, textColor=ACCENT,
        leading=16, spaceAfter=4)

    bullet_s = ParagraphStyle("Bullet", parent=body_s,
        leftIndent=22, firstLineIndent=-10, spaceAfter=3)

    subbullet_s = ParagraphStyle("SubBullet", parent=body_s,
        leftIndent=36, firstLineIndent=-10, spaceAfter=2,
        fontSize=9.5, textColor=MUTED)

    num_s = ParagraphStyle("Num", parent=body_s,
        leftIndent=22, firstLineIndent=-16, spaceAfter=3)

    muted_s = ParagraphStyle("Muted", parent=body_s,
        textColor=MUTED, fontName="Helvetica-Oblique", fontSize=10)

    small_s = ParagraphStyle("Small", parent=body_s,
        fontSize=9, textColor=MUTED)

    footer_s = ParagraphStyle("Footer", parent=small_s, alignment=1, spaceAfter=0)

    story = []

    # Title
    story.append(Paragraph(notes_dict.get("title", "Lecture Notes"), title_s))
    story.append(HRFlowable(width="100%", thickness=2, color=BRAND, spaceAfter=6))

    # Stats row
    wc   = notes_dict.get("word_count", 0)
    nsec = len(notes_dict.get("sections", []))
    nkp  = len(notes_dict.get("key_points", []))
    story.append(Paragraph(
        f"<b>Words:</b> {wc:,} &nbsp;&nbsp; <b>Key Concepts:</b> {nkp} &nbsp;&nbsp; <b>Sections:</b> {nsec}",
        small_s
    ))
    story.append(Spacer(1, 0.4*cm))

    # Overview
    summary = _strip_md(notes_dict.get("summary", ""))
    if summary:
        story.append(Paragraph("Overview", h1_s))
        story.append(Paragraph(summary, muted_s))
        story.append(Spacer(1, 0.3*cm))

    # Key Concepts — numbered list
    key_points = notes_dict.get("key_points", [])
    if key_points:
        story.append(Paragraph("Key Concepts", h1_s))
        items = []
        for i, pt in enumerate(key_points, 1):
            clean = _strip_md(pt)
            if clean:
                items.append(ListItem(
                    Paragraph(clean, body_s),
                    value=i, bulletColor=BRAND, leftIndent=18
                ))
        if items:
            story.append(ListFlowable(items, bulletType='1', leftIndent=16, bulletFormat="%s."))
        story.append(Spacer(1, 0.35*cm))

    # Detailed Sections
    sections = notes_dict.get("sections", [])
    if sections:
        story.append(Paragraph("Detailed Notes", h1_s))
        for sec in sections:
            heading = _strip_md(sec.get("heading", ""))
            if heading:
                story.append(Paragraph(heading, h2_s))
            content = _strip_md(sec.get("content", ""))
            for line in content.split('\n'):
                raw = line
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.1*cm))
                    continue
                if raw.startswith('  - ') or raw.startswith('  • '):
                    # Sub-bullet — deeper indent, muted colour
                    story.append(Paragraph(
                        f"&#9702; {line.lstrip('-• ')}",
                        subbullet_s
                    ))
                elif re.match(r'^\s{2,}\d+[.)]', raw):
                    # Sub-numbered — deeper indent, muted
                    story.append(Paragraph(
                        re.sub(r'^\s+', '', raw),
                        subbullet_s
                    ))
                elif line.startswith('• '):
                    story.append(Paragraph(f"• {line[2:]}", bullet_s))
                elif re.match(r'^\d+[.)]\s', line):
                    story.append(Paragraph(line, num_s))
                else:
                    story.append(Paragraph(line, body_s))
        story.append(Spacer(1, 0.4*cm))

    # Full Transcript
    full_tr = notes_dict.get("full_transcript", "")
    if full_tr:
        story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=4))
        story.append(Paragraph("Full Transcript (English)", h1_s))
        for chunk in (full_tr[i:i+600] for i in range(0, len(full_tr), 600)):
            story.append(Paragraph(chunk.replace('\n', '<br/>'), small_s))

    # Footer
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
    story.append(Paragraph("Generated by AudioNotes AI", footer_s))

    doc.build(story)
    buf.seek(0)
    return buf.read()
